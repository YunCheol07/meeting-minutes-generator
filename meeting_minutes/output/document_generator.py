"""Word 문서 생성 모듈 - 회의록을 Word 파일로 출력"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from ..core.state_schema import MeetingState


class MeetingMinutesDocGenerator:
    """회의록 Word 문서 생성기
    
    추출된 회의 정보를 구조화된 Word 문서로 생성합니다.
    """
    
    def __init__(self):
        """문서 생성기 초기화"""
        self.doc = Document()
        self._setup_styles()
    
    def _setup_styles(self):
        """문서 스타일 설정 - 한글 폰트 지원"""
        # Normal 스타일 설정
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Malgun Gothic'  # 맑은 고딕
        font.size = Pt(10)
        
        # 한글 폰트 명시적 설정 (동아시아 문자)
        style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    
    def _add_title(self, title: str):
        """문서 제목 추가
        
        Args:
            title: 회의 제목
        """
        heading = self.doc.add_heading(title, level=0)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # 제목 폰트 설정
        for run in heading.runs:
            run.font.name = 'Malgun Gothic'
            run.font.size = Pt(18)
            run.font.bold = True
            # 한글 폰트 명시적 설정
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    
    def _add_section(self, title: str, level: int = 1):
        """섹션 제목 추가
        
        Args:
            title: 섹션 제목
            level: 제목 레벨 (1-3)
        """
        heading = self.doc.add_heading(title, level=level)
        
        # 섹션 제목 폰트 설정
        for run in heading.runs:
            run.font.name = 'Malgun Gothic'
            run.font.color.rgb = RGBColor(0, 51, 102)  # 네이비 색상
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    
    def _add_paragraph(self, text: str, bold: bool = False):
        """일반 문단 추가
        
        Args:
            title: 문단 텍스트
            bold: 볼드 여부
        
        Returns:
            Paragraph: 추가된 문단 객체
        """
        para = self.doc.add_paragraph(text)
        
        # 폰트 설정
        for run in para.runs:
            run.font.name = 'Malgun Gothic'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
            if bold:
                run.bold = True
        
        return para
    
    def _add_bullet_list(self, items: list):
        """불릿 리스트 추가
        
        Args:
            items: 리스트 항목들
        """
        for item in items:
            para = self.doc.add_paragraph(str(item), style='List Bullet')
            # 폰트 설정
            for run in para.runs:
                run.font.name = 'Malgun Gothic'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    
    def _add_table(self, headers: list, rows: list):
        """테이블 추가
        
        Args:
            headers: 헤더 행
            rows: 데이터 행들
        
        Returns:
            Table: 추가된 테이블 객체
        """
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.style = 'Light Grid Accent 1'
        
        # 헤더 행 추가
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            # 헤더 폰트 설정
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Malgun Gothic'
                    run.font.bold = True
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
        
        # 데이터 행 추가
        for row_data in rows:
            row_cells = table.add_row().cells
            for i, cell_data in enumerate(row_data):
                row_cells[i].text = str(cell_data)
                # 셀 폰트 설정
                for paragraph in row_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Malgun Gothic'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
        
        return table
    
    def generate(self, state: MeetingState, output_path: str) -> str:
        """회의록 문서 생성
        
        Args:
            state: 최종 상태 (모든 정보 포함)
            output_path: 출력 파일 경로
        
        Returns:
            str: 생성된 파일 경로
        """
        print("\n[Step 8/8] Word 문서 생성 중...")
        
        # 1. 제목
        self._add_title(state["meeting_title"])
        
        # 2. 기본 정보
        self._add_section("회의 정보")
        self._add_paragraph(f"일시: {state['meeting_date']}")
        self._add_paragraph(f"참석자: {', '.join(state['participants'])}")
        self.doc.add_paragraph()  # 빈 줄
        
        # 3. 회의 요약
        self._add_section("회의 요약")
        self._add_paragraph(state["summary"])
        self.doc.add_paragraph()
        
        # 4. 안건
        if state["agenda_items"] and state["agenda_items"] != ["안건 내용 없음"]:
            self._add_section("안건")
            self._add_bullet_list(state["agenda_items"])
            self.doc.add_paragraph()
        
        # 5. 논의 내용
        if state["discussions"]:
            self._add_section("논의 내용")
            for disc in state["discussions"]:
                if disc["topic"] != "논의 내용":  # 에러 메시지 제외
                    self._add_paragraph(f"• {disc['topic']}", bold=True)
                    self._add_paragraph(f"  {disc['content']}")
            self.doc.add_paragraph()
        
        # 6. 결정 사항
        if state["decisions"] and state["decisions"] != ["특별한 결정 사항 없음"]:
            self._add_section("결정 사항")
            self._add_bullet_list(state["decisions"])
            self.doc.add_paragraph()
        
        # 7. 액션 아이템
        if state["action_items"]:
            # "후속 조치 없음"이 아닌 경우만 표시
            real_actions = [
                item for item in state["action_items"]
                if item["task"] != "후속 조치 없음"
            ]
            
            if real_actions:
                self._add_section("액션 아이템")
                headers = ["작업 내용", "담당자", "마감일"]
                rows = [
                    [item["task"], item["assignee"], item["deadline"]]
                    for item in real_actions
                ]
                self._add_table(headers, rows)
        
        # 8. 문서 저장
        self.doc.save(output_path)
        print(f"✓ 회의록 생성 완료: {output_path}")
        
        return output_path
